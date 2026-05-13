import io, json, logging
from pathlib import Path
from fastapi import APIRouter, HTTPException, UploadFile, File, Request
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from app.analyzers.txt_analyzer import analyze_txt_prompts
from app.analyzers.word_to_html import convert_word_to_html
from app.rag.rag_service import RAGService

router = APIRouter()
logger = logging.getLogger(__name__)


def get_settings(request: Request):
    return request.app.state.settings


class ChatRequest(BaseModel):
    message: str
    context: str = ""
    user_id: int = 0


class CVAnalyzeRequest(BaseModel):
    cv_text: str


@router.post("/analyze-file")
async def analyze_file(request: Request, file: UploadFile = File(...)):
    settings = get_settings(request)
    content = await file.read()
    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in ["pdf", "docx", "txt"]:
        raise HTTPException(status_code=400, detail=f"AI analiz için desteklenmeyen dosya tipi: {ext}")

    rag = RAGService(settings.GEMINI_API_KEY, settings.GEMINI_MODEL)
    result = await rag.analyze_cv_file(content, ext)
    return result


@router.post("/analyze-file-by-id/{file_id}")
async def analyze_file_by_id(file_id: int, request: Request):
    """Önceden yüklenmiş bir dosyayı dosya ID'siyle analiz eder"""
    import time
    pool = request.app.state.db_pool
    settings = get_settings(request)

    async with pool.acquire() as conn:
        row = await conn.fetchrow("SELECT user_id, value FROM files WHERE id = $1", file_id)
    if not row:
        raise HTTPException(status_code=404, detail="File not found")

    user_id = row["user_id"]
    value = json.loads(row["value"])
    path = Path(value.get("path", ""))
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    content = path.read_bytes()
    ext = value.get("extension", "").lower()

    if ext not in ["pdf", "docx", "txt"]:
        raise HTTPException(status_code=400, detail=f"AI analiz için desteklenmeyen dosya tipi: {ext}")

    rag = RAGService(settings.GEMINI_API_KEY, settings.GEMINI_MODEL)
    analysis_res = await rag.analyze_cv_file(content, ext)

    if not analysis_res.get("success"):
        raise HTTPException(status_code=500, detail=f"AI Analiz Hatası: {analysis_res.get('error')}")

    cv_data = analysis_res.get("data", {})

    # Update file record with AI result
    async with pool.acquire() as conn:
        await conn.execute(
            "UPDATE files SET status='analyzed', value = value || $1::jsonb WHERE id = $2",
            json.dumps({"ai_analyzed": True, "ai_result": cv_data}), file_id
        )

        # Retrieve and auto-populate student record
        student_row = await conn.fetchrow("SELECT id, value FROM students WHERE user_id = $1", user_id)
        if student_row:
            student_id = student_row["id"]
            student_value = json.loads(student_row["value"]) if student_row["value"] else {}

            student_value["cv_url"] = f"/api/v1/files/download/{file_id}"

            # Safely merge skills
            if not student_value.get("skills"):
                student_value["skills"] = cv_data.get("skills", [])
            # Safely merge languages
            if not student_value.get("languages"):
                student_value["languages"] = cv_data.get("languages", [])
            # Safely set bio
            if not student_value.get("bio"):
                student_value["bio"] = cv_data.get("summary", "")

            # Safely map and merge projects
            if not student_value.get("projects"):
                projects_mapped = []
                current_time = int(time.time() * 1000)
                for idx, p in enumerate(cv_data.get("projects", [])):
                    techs = ", ".join(p.get("technologies", [])) if isinstance(p.get("technologies"), list) else ""
                    projects_mapped.append({
                        "id": current_time + idx,
                        "title": p.get("name", "Proje"),
                        "description": p.get("description", ""),
                        "tech": techs
                    })
                student_value["projects"] = projects_mapped

            await conn.execute(
                "UPDATE students SET value = $1::jsonb, updated_at = NOW() WHERE id = $2",
                json.dumps(student_value), student_id
            )

    return cv_data


@router.post("/chat")
async def ai_chat(body: ChatRequest, request: Request):
    settings = get_settings(request)
    rag = RAGService(settings.GEMINI_API_KEY, settings.GEMINI_MODEL)
    response = await rag.chat(body.message, body.context)
    return {"response": response, "model": settings.GEMINI_MODEL}


@router.post("/analyze-cv")
async def analyze_cv(body: CVAnalyzeRequest, request: Request):
    settings = get_settings(request)
    rag = RAGService(settings.GEMINI_API_KEY, settings.GEMINI_MODEL)
    result = await rag.analyze_cv(body.cv_text)
    return result


@router.post("/search-students")
async def search_students(request: Request, query: str, user_id: int):
    settings = get_settings(request)
    pool = request.app.state.db_pool
    rag = RAGService(settings.GEMINI_API_KEY, settings.GEMINI_MODEL)

    async with pool.acquire() as conn:
        rows = await conn.fetch(
            "SELECT st.id, u.value as user_val, st.value as student_val FROM students st JOIN users u ON u.id=st.user_id WHERE u.is_active=TRUE LIMIT 50"
        )

    profiles = []
    for row in rows:
        uv = json.loads(row["user_val"])
        sv = json.loads(row["student_val"])
        profiles.append({
            "id": row["id"],
            "name": f"{uv.get('name','')} {uv.get('surname','')}",
            "skills": sv.get("skills", []),
            "gpa": sv.get("gpa"),
            "department": sv.get("department", ""),
        })

    matches = await rag.search_students(query, profiles)
    return {"query": query, "matches": matches, "total": len(matches)}
@router.post("/convert-word-to-html")
async def convert_word(request: Request, file: UploadFile = File(...)):
    settings = get_settings(request)
    content = await file.read()
    result = await convert_word_to_html(content, settings.GEMINI_API_KEY)
    if not result.get("success"):
        raise HTTPException(status_code=500, detail=result.get("error"))
    return result


@router.post("/analyze-txt")
async def analyze_txt(request: Request, file: UploadFile = File(...)):
    settings = get_settings(request)
    content = (await file.read()).decode("utf-8", errors="ignore")
    result = await analyze_txt_prompts(content, settings.GEMINI_API_KEY)
    return result


@router.post("/convert-file-by-id/{file_id}")
async def convert_file_by_id(file_id: int, request: Request, target_format: str = "html"):
    """Önceden yüklenmiş bir dosyayı farklı bir formata (örn: html) dönüştürür"""
    pool = request.app.state.db_pool
    settings = get_settings(request)

    async with pool.acquire() as conn:
        row = await conn.fetchrow("SELECT user_id, value FROM files WHERE id = $1", file_id)
    if not row:
        raise HTTPException(status_code=404, detail="File not found")

    value = json.loads(row["value"])
    path = Path(value.get("path", ""))
    if not path.exists():
        raise HTTPException(status_code=404, detail="File not found on disk")

    content = path.read_bytes()
    ext = value.get("extension", "").lower()

    if target_format == "html" and ext == "docx":
        result = await convert_word_to_html(content, settings.GEMINI_API_KEY)
    elif target_format == "analysis" and ext == "txt":
        txt_content = content.decode("utf-8", errors="ignore")
        result = await analyze_txt_prompts(txt_content, settings.GEMINI_API_KEY)
        result["success"] = True # analyze_txt_prompts returns a dict without success key usually, adding it for consistency
    else:
        raise HTTPException(status_code=400, detail=f"{ext} dosyasından {target_format} formatına dönüşüm desteklenmiyor.")

    if not result.get("success"):
        raise HTTPException(status_code=500, detail=f"Dönüşüm Hatası: {result.get('error')}")

    # Sonucu veritabanına kaydet
    async with pool.acquire() as conn:
        update_data = {"ai_converted": True, "target_format": target_format}
        if target_format == "html":
            update_data["converted_html"] = result.get("html")
        else:
            update_data["ai_result"] = result.get("results")
            
        await conn.execute(
            "UPDATE files SET status='analyzed', value = value || $1::jsonb WHERE id = $2",
            json.dumps(update_data), file_id
        )

    return result


# ─────────────────────────────────────────────────────────────────────────────
# FORMAT CONVERSION: txt / pdf / docx / xlsx  ↔  each other
# ─────────────────────────────────────────────────────────────────────────────

SUPPORTED_FORMATS = {"txt", "pdf", "docx", "xlsx"}

MIME_TYPES = {
    "txt":  "text/plain",
    "pdf":  "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
}


async def _read_source_as_text(content: bytes, ext: str) -> str:
    """Any supported source format → plain text"""
    if ext == "txt":
        return content.decode("utf-8", errors="ignore")
    if ext == "pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"PDF okunamadı: {e}")
    if ext == "docx":
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"DOCX okunamadı: {e}")
    if ext == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.load_workbook(io.BytesIO(content), data_only=True)
            lines = []
            for ws in wb.worksheets:
                lines.append(f"=== {ws.title} ===")
                for row in ws.iter_rows(values_only=True):
                    lines.append("\t".join(str(c) if c is not None else "" for c in row))
            return "\n".join(lines)
        except Exception as e:
            raise HTTPException(status_code=422, detail=f"XLSX okunamadı: {e}")
    raise HTTPException(status_code=400, detail=f"Desteklenmeyen kaynak format: {ext}")


async def _build_target(text: str, target_ext: str, api_key: str) -> bytes:
    """Plain text -> target format bytes"""
    if target_ext == "txt":
        return text.encode("utf-8")

    if target_ext == "pdf":
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.lib.units import cm
            buf = io.BytesIO()
            doc = SimpleDocTemplate(buf, pagesize=A4, leftMargin=2*cm, rightMargin=2*cm,
                                    topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            story = []
            for line in text.splitlines():
                if line.strip():
                    safe = line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
                    story.append(Paragraph(safe, styles["Normal"]))
                    story.append(Spacer(1, 6))
            doc.build(story)
            return buf.getvalue()
        except ImportError:
            logger.warning("reportlab bulunamadi, PDF icin TXT fallback kullaniliyor")
            raise HTTPException(
                status_code=503,
                detail="PDF üretimi için 'reportlab' paketi gerekli. Lütfen ai-service container'ini yeniden build edin."
            )

    if target_ext == "docx":
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc()
            for line in text.splitlines():
                doc.add_paragraph(line)
            buf = io.BytesIO()
            doc.save(buf)
            return buf.getvalue()
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"DOCX üretime hatası: {e}")

    if target_ext == "xlsx":
        try:
            import openpyxl
            wb = openpyxl.Workbook()
            ws = wb.active
            ws.title = "CV"
            for i, line in enumerate(text.splitlines(), start=1):
                ws.cell(row=i, column=1, value=line)
            buf = io.BytesIO()
            wb.save(buf)
            return buf.getvalue()
        except ImportError:
            raise HTTPException(
                status_code=503,
                detail="Excel üretimi için 'openpyxl' paketi gerekli. Lütfen ai-service container'ini yeniden build edin."
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Excel üretim hatası: {e}")

    raise HTTPException(status_code=400, detail=f"Desteklenmeyen hedef format: {target_ext}")


@router.post("/convert-format")
async def convert_format(
    request: Request,
    file: UploadFile = File(...),
    target_format: str = "txt",
):
    """
    Dosyayı desteklenen formatlar arasında dönüştürür ve indirmek için döndürür.
    Desteklenen: txt ↔ pdf ↔ docx ↔ xlsx
    """
    settings = get_settings(request)

    filename = file.filename or "file"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext not in SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Kaynak format desteklenmiyor: .{ext}. Desteklenenler: {', '.join(SUPPORTED_FORMATS)}")
    if target_format not in SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Hedef format desteklenmiyor: .{target_format}. Desteklenenler: {', '.join(SUPPORTED_FORMATS)}")
    if ext == target_format:
        raise HTTPException(status_code=400, detail="Kaynak ve hedef format aynı olamaz.")

    content = await file.read()

    # Step 1: Extract text from source
    source_text = await _read_source_as_text(content, ext)

    # Step 2: Build target file bytes
    result_bytes = await _build_target(source_text, target_format, settings.GEMINI_API_KEY)

    import urllib.parse
    base_name = filename.rsplit(".", 1)[0]
    download_filename = f"{base_name}_converted.{target_format}"
    encoded_filename = urllib.parse.quote(download_filename)
    mime = MIME_TYPES.get(target_format, "application/octet-stream")

    return StreamingResponse(
        io.BytesIO(result_bytes),
        media_type=mime,
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"},
    )


@router.post("/export-cv")
async def export_cv(
    request: Request,
    target_format: str = "txt",
):
    """
    İstek gövdesindeki CV JSON verisini seçilen formata (txt/pdf/docx/xlsx) çevirir.
    """
    settings = get_settings(request)
    body = await request.json()
    cv_data: dict = body.get("cv_data", {})

    if target_format not in SUPPORTED_FORMATS:
        raise HTTPException(status_code=400, detail=f"Desteklenmeyen format: {target_format}")

    # Build plain-text representation of CV
    lines = []
    if cv_data.get("name"):
        lines += [cv_data["name"], "=" * len(cv_data["name"]), ""]
    if cv_data.get("summary"):
        lines += ["ÖZET", "-" * 10, cv_data["summary"], ""]
    if cv_data.get("skills"):
        lines += ["YETENEKler", "-" * 10] + [f"• {s}" for s in cv_data["skills"]] + [""]
    if cv_data.get("languages"):
        lines += ["DİLLER", "-" * 10] + [f"• {l}" for l in cv_data["languages"]] + [""]
    if cv_data.get("education"):
        lines += ["EĞİTİM", "-" * 10]
        for e in cv_data["education"]:
            lines.append(f"{e.get('degree', '')} — {e.get('school', '')} ({e.get('year', '')})")
        lines.append("")
    if cv_data.get("experience"):
        lines += ["İŞ DENEYİMİ", "-" * 10]
        for e in cv_data["experience"]:
            lines.append(f"{e.get('title', '')} @ {e.get('company', '')} — {e.get('duration', '')}")
        lines.append("")
    if cv_data.get("projects"):
        lines += ["PROJELER", "-" * 10]
        for p in cv_data["projects"]:
            techs = ", ".join(p.get("technologies", [])) if isinstance(p.get("technologies"), list) else ""
            lines.append(f"{p.get('name', '')} — {p.get('description', '')} [{techs}]")
        lines.append("")
    if cv_data.get("improvements"):
        lines += ["AI İYİLEŞTİRME ÖNERİLERİ", "-" * 10] + [f"✓ {s}" for s in cv_data["improvements"]] + [""]

    text = "\n".join(lines)
    result_bytes = await _build_target(text, target_format, settings.GEMINI_API_KEY)

    import urllib.parse
    mime = MIME_TYPES.get(target_format, "application/octet-stream")
    name = cv_data.get("name", "CV").replace(" ", "_")
    download_filename = f"{name}_CV.{target_format}"
    encoded_filename = urllib.parse.quote(download_filename)

    return StreamingResponse(
        io.BytesIO(result_bytes),
        media_type=mime,
        headers={"Content-Disposition": f"attachment; filename*=utf-8''{encoded_filename}"},
    )
